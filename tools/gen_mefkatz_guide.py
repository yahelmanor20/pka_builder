# -*- coding: utf-8 -*-
"""
מחולל "מדריך למפקץ — איך לכתוב לו"ז נכון ביומן" (DOCX עברית RTL, מעוצב).
התוכן נגזר ישירות מהפרסר ב-index.html (COLOR_PLATOON, parseDescription, computeIssues,
detectPlatoon, זיהוי יום מפוצל ופעילויות מקבילות) כך שכל כלל תואם להתנהגות הכלי בפועל.
עיצוב מקצועי: באנר כותרת · כרטיסי כללים עם פס מספור צבעוני · תיבות נכון/שגוי · RTL.
פלט: מדריך_כתיבת_לוז_ביומן.docx בשורש הפרויקט.
"""
import os
from docx.shared import Pt, RGBColor, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# reuse של מנגנון ה-RTL/עיצוב מהמחולל הקיים
from gen_docx import (
    base_doc, _el, _run, _prtl, _cell, _vcenter, _full_width, _shade,
)

HERE = os.path.dirname(__file__)
OUT = os.path.join(HERE, "..", "מדריך_כתיבת_לוז_ביומן.docx")

# פלטת צבעים
BLUE = RGBColor(0x1F, 0x4E, 0x79)          # כחול מותג
BLUE_HEX = "1F4E79"
BAND = "DEEAF6"                            # פס כותרת תכלת
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x33, 0x33, 0x33)
GREEN = RGBColor(0x37, 0x6E, 0x2A); GREEN_FILL = "E2EFDA"
RED = RGBColor(0xA0, 0x00, 0x00); RED_FILL = "FBE3DD"
GRAY = RGBColor(0x70, 0x70, 0x70)
SOFT = "C9C9C9"                           # קו טבלה עדין

# צבע ביומן (hex משוער של Google Calendar) → שם בבורר → משמעות בכלי
COLORS = [
    ("D50000", WHITE, "עגבנייה (Tomato)",    "פלוגה א׳"),
    ("F4511E", WHITE, "מנדרינה (Tangerine)",  "פלוגה ב׳"),
    ("F6BF26", DARK,  "בננה (Banana)",        "פלוגה ג׳"),
    ("0B8043", WHITE, "בזיליקום (Basil)",      "פלוגה ד׳"),
    ("E67C73", DARK,  "פלמינגו (Flamingo)",    "לו״ז מקביל — הסמכות / רענונים / חובשים / תאג״ד → נספח"),
    ("616161", WHITE, "גרפיט (Graphite)",      "פלס״ם → נספח פלס״ם"),
]


# ---------- כלי עיצוב ----------
def _set_borders(t, color="FFFFFF", sz="4", val="single"):
    b = _el("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b.append(_el("w:" + e, **{"w:val": val, "w:sz": sz, "w:space": "0", "w:color": color}))
    t._tbl.tblPr.append(b)


def _no_space(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def spacer(doc, pts=4):
    p = doc.add_paragraph(); _prtl(p); _no_space(p); p.paragraph_format.space_after = Pt(pts)
    _run(p.add_run(""), size=2)


def para(doc, runs, size=10.5, after=2, indent=0):
    """runs = רשימת (טקסט, bold, color)."""
    p = doc.add_paragraph(); _prtl(p); _no_space(p, after=after)
    if indent:
        p.paragraph_format.right_indent = Twips(indent)
    for txt, bold, color in runs:
        _run(p.add_run(txt), size=size, bold=bold, color=color)
    return p


def bullet(doc, text, size=10.5, bold_head=None):
    p = doc.add_paragraph(); _prtl(p); _no_space(p, after=2)
    p.paragraph_format.right_indent = Twips(260)
    if bold_head:
        _run(p.add_run("◆  "), size=size, color=BLUE)
        _run(p.add_run(bold_head), size=size, bold=True, color=DARK)
        _run(p.add_run(text), size=size, color=DARK)
    else:
        _run(p.add_run("◆  "), size=size, color=BLUE)
        _run(p.add_run(text), size=size, color=DARK)
    return p


def banner(doc):
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t._tbl.tblPr.append(_el("w:bidiVisual")); _set_borders(t, val="nil")
    _cell(t.rows[0].cells[0], "איך לכתוב לו״ז נכון ביומן", bold=True, size=20, fill=BLUE_HEX,
          color=WHITE, center=True); _vcenter(t.rows[0].cells[0])
    _cell(t.rows[1].cells[0], "מדריך למפקץ · בא״פ 8222 · ביה״ס ללחימה", bold=False, size=11,
          fill=BLUE_HEX, color=WHITE, center=True)
    _full_width(t, [1])
    return t


def section(doc, num, head):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t._tbl.tblPr.append(_el("w:bidiVisual")); _set_borders(t, val="nil")
    _cell(t.rows[0].cells[0], str(num), bold=True, size=15, fill=BLUE_HEX, color=WHITE, center=True)
    _vcenter(t.rows[0].cells[0])
    _cell(t.rows[0].cells[1], head, bold=True, size=12.5, fill=BAND, color=BLUE)
    _vcenter(t.rows[0].cells[1])
    _full_width(t, [0.6, 9.4])
    return t


def good_bad(doc, good, bad):
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t._tbl.tblPr.append(_el("w:bidiVisual")); _set_borders(t, color=SOFT)
    _cell(t.rows[0].cells[0], "✓  נכון", bold=True, size=10, fill=GREEN_FILL, color=GREEN, center=True)
    _cell(t.rows[0].cells[1], "✗  שגוי", bold=True, size=10, fill=RED_FILL, color=RED, center=True)
    _cell(t.rows[1].cells[0], good, size=10, color=DARK); _vcenter(t.rows[1].cells[0])
    _cell(t.rows[1].cells[1], bad, size=10, color=DARK); _vcenter(t.rows[1].cells[1])
    _full_width(t, [1, 1])
    return t


def color_table(doc):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t._tbl.tblPr.append(_el("w:bidiVisual")); _set_borders(t, color="FFFFFF", sz="8")
    _cell(t.rows[0].cells[0], "צבע האירוע ביומן", bold=True, size=10.5, fill=BLUE_HEX,
          color=WHITE, center=True); _vcenter(t.rows[0].cells[0])
    _cell(t.rows[0].cells[1], "כך הכלי משייך אותו", bold=True, size=10.5, fill=BLUE_HEX,
          color=WHITE, center=True); _vcenter(t.rows[0].cells[1])
    for hexfill, txtcolor, name, meaning in COLORS:
        cells = t.add_row().cells
        _cell(cells[0], name, bold=True, size=11, fill=hexfill, color=txtcolor, center=True)
        _vcenter(cells[0])
        _cell(cells[1], meaning, size=10.5, color=DARK); _vcenter(cells[1])
    _full_width(t, [3.0, 5.0])
    return t


def callout(doc, text, fill, txtcolor):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t._tbl.tblPr.append(_el("w:bidiVisual")); _set_borders(t, val="nil")
    _cell(t.rows[0].cells[0], text, bold=True, size=10.5, fill=fill, color=txtcolor, center=True)
    _vcenter(t.rows[0].cells[0])
    _full_width(t, [1])
    return t


# ---------- בניית המסמך ----------
def build(path):
    doc = base_doc()
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(doc.sections[0], m, Inches(0.55))

    banner(doc)
    para(doc, [("הכלי בונה את הפק״א והלו״ז ", False, DARK),
               ("אוטומטית מתוך היומן", True, DARK),
               (". אם תכתוב את האירועים לפי 8 הכללים שכאן — הוא יקרא אותם נכון, "
                "וכמעט לא תצטרך לתקן ידנית. כל כלל כאן תואם בדיוק לאופן שבו הכלי מפענח את היומן.",
                False, DARK)], size=10.5, after=4)

    # 1 — צבע
    section(doc, 1, "הכי חשוב: צבע האירוע קובע את הפלוגה")
    callout(doc, "הצבע גובר על הכותרת. אירוע בלי צבע — לא משויך לאף פלוגה!", RED_FILL, RED)
    spacer(doc, 2)
    color_table(doc)
    bullet(doc, "הכלי בודק קודם את הצבע, ורק אם אין צבע — מנסה להבין פלוגה מהכותרת.", bold_head="למה זה ראשון:  ")
    bullet(doc, "ביומן: פותחים את האירוע ← לוחצים על עיגול הצבע ← בוחרים מהרשימה את הצבע הנכון.", bold_head="איך צובעים:  ")
    bullet(doc, "אם הכותרת אומרת פלוגה אחת אבל הצבע אחר — הכלי הולך לפי הצבע ורושם הערה בלו״ז.", bold_head="סתירה:  ")

    # 2 — כותרת
    section(doc, 2, "כותרת אירוע פלוגתי")
    para(doc, [("פורמט: ", True, DARK), ("פלוגה א׳ - שם הפעילות", True, BLUE),
               ("   ←   הטקסט שאחרי המקף הוא שם הפעילות שיופיע בלו״ז.", False, DARK)])
    good_bad(doc, "פלוגה א׳ - מטווח קלעים", "מטווח קלעים (בלי ציון פלוגה / בלי צבע)")

    # 3 — שעה
    section(doc, 3, "שעה — בשדה השעה של האירוע, לא בכותרת")
    bullet(doc, "מלא שעת התחלה וסיום אמיתיות באירוע. הכלי קורא אותן ובונה מהן את עמודת השעה.")
    bullet(doc, "אירוע ״כל היום״ או בלי שעה מוגדרת → מסומן באזהרה ״ללא שעה״.")
    good_bad(doc, "אירוע מתוזמן 08:00–11:00", "אירוע ״כל היום״, או ״08:00 מטווח״ ככותרת")

    # 4 — מיקום
    section(doc, 4, "מיקום")
    bullet(doc, "מלא את שדה המיקום של האירוע ביומן.")
    bullet(doc, "שדה ריק → הכלי כותב ״מיקום לא ידוע״ ומסמן באזהרה צהובה.")
    bullet(doc, "אפשר גם דרך התיאור: שורת  ", bold_head=None)
    para(doc, [("        מיקום: <שם המקום>", True, BLUE)], size=10)
    good_bad(doc, "שדה מיקום: ״שטח 9א׳״   ·   או בתיאור: מיקום: יובלים",
             "שדה מיקום ריק והמקום כתוב רק בתוך הכותרת")

    # 5 — תיאור
    section(doc, 5, "תיאור — שדות מסודרים")
    para(doc, [("כל שדה בשורה נפרדת, עם נקודתיים. הכלי מזהה בדיוק את שלוש המילים האלה:", False, DARK)])
    bullet(doc, "<שם המעביר>  — חסר → אזהרה ״חסר גורם מעביר״.", bold_head="גורם מעביר: ")
    bullet(doc, "<טקסט חופשי>", bold_head="הערות: ")
    bullet(doc, "<מקום> — דורס את שדה המיקום.", bold_head="מיקום: ")
    bullet(doc, "כל שורה שאינה באחד הפורמטים האלה נדבקת אוטומטית לשדה ההערות.", bold_head="שאר השורות:  ")
    bullet(doc, "באירוע טכניקות/תרגול — פירוט בתיאור נדבק בסוגריים לשם הפעילות.", bold_head="טכניקות:  ")

    # 6 — מקבילות
    section(doc, 6, "פעילויות מקבילות (לא בלו״ז היומי)")
    bullet(doc, "חובשים · תאג״ד · סמב״צים · הסמכות · רענונים — צבע פלמינגו → עוברות לנספח המקביל.")
    bullet(doc, "פלס״ם — צבע גרפיט → נספח פלס״ם נפרד.")
    callout(doc, "חריג חשוב: ״מרחב חכם״ תמיד נשאר בלו״ז היומי — גם אם צבעת אותו פלמינגו.", "FFF2CC", RGBColor(0x80, 0x60, 0x00))

    # 7 — יום מפוצל
    section(doc, 7, "יום מפוצל (2+ פלוגות באותו יום)")
    bullet(doc, "כשהכלי מזהה 2 פלוגות שונות או יותר באותו יום — הטבלה הופכת אוטומטית לעמודות פלוגה.")
    bullet(doc, "פעילות לכל הגדוד ביום מפוצל — השאר אותה בלי צבע פלוגה, כדי שתופיע כשורה רוחבית אחת ולא בעמודה אחת.")

    # 8 — צהוב
    section(doc, 8, "לפני שסיימת — בדוק את הצהוב")
    bullet(doc, "כל אזהרה צהובה בתצוגה = משהו לתקן ביומן: ", bold_head=None)
    para(doc, [("        ללא שעה   ·   מיקום לא ידוע   ·   חסר גורם מעביר", True, RED)], size=10.5)

    # אל תעשה
    spacer(doc, 3)
    callout(doc,
            "✗ אל תעשה:   שעה בכותרת במקום בשדה   ·   אירוע בלי צבע   ·   מיקום ריק   ·   "
            "שיוך פלוגה לפי כותרת בלבד   ·   פעילות מקבילה בצבע פלוגה",
            RED_FILL, RED)

    doc.save(path)
    print("נשמר:", os.path.abspath(path))


if __name__ == "__main__":
    build(OUT)
